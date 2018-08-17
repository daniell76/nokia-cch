from cch_doc_xls import *

from docx import Document
from docx.shared import *
from docx.enum.text import *
from docx.enum.style import WD_STYLE_TYPE
from copy import deepcopy

#from docx.shared import Inches
#from docx.shared import Pt
import os
import re
import collections
import xml.etree.ElementTree as ET
import datetime

class CCH_DOCX(object):
  def __init__(self,filename=None):
    if filename is None:
      self.document = Document()
    else:
      try:
        self.document = Document(filename)
      except:
        self.document = None
#    self.__headings = collections.OrderedDict()
#    self.__contents = collections.OrderedDict()
    self.param = {}
    self.contents = {}
    self.paragraphs = collections.OrderedDict()
    self.table = {}
    self.tables = 0
    self.images = 0
    self.tagre = re.compile('([^<>]*)(<[^<>]*>){0,1}')
  
  def loadParameters(self,filename=None):
    assert filename is not None
    xls = CCH_PARAM_XLS(filename)
    xls.parse()
    self.param = deepcopy(xls.param)
    try:
      for t in xls.customTable:
        self.table[t] = deepcopy(xls.customTable[t])
    except:
      pass
    
  def loadContents(self,filename=None):
    assert filename is not None
    xls = CCH_CONTENT_XLS(filename)
    xls.parse()
    self.contents = deepcopy(xls.contents)
    try:
      for t in xls.customTable:
        self.table[t] = {}
        for k in ['title','body']:
          self.table[t][k] = deepcopy(xls.customTable[t][k])
    except Exception as e:
      print(str(e))
      pass

  def applyParam(self, param_name):
    try:
      l = param_name.split('/')
      return self.param[l[0]][l[1]]
    except:
      return None
  
  def parseText(self,txt=None):
    # only substitute parameters.
    if txt is None: return None
    if len(txt) == 0: return None
    tokens = []
    m = self.tagre.findall(txt)
    for l in m:
      for i in range(2):
        try:
          assert len(l[i]) > 0
          tokens.append(l[i])
        except:
          pass
    ret = ''
    for token in tokens:
      tag = CCH_DOCX.parseTags(token)
      if tag['type'] == 'param':
        ret += str(self.applyParam(tag['tagName']))
      else:
        ret += token
    return ret
        
  def parseTables(self):
    for k in self.table:
      # Title cannot be parameterised
      for r in self.table[k]['body']:
        for c in r:
          r[c] = self.parseText(r[c])
  
  def addTable(self,tableName,caption=None,afterParagraph=None):
    try:
      table = self.table[tableName]
    except:
      return None
    rows = len(table['body']) + 1
    cols = len(table['title'])
    t = self.document.add_table(rows=rows+1, cols=cols)
    t.style = 'Table Grid'
    #CCH_DOCX.set_table_font_size(t,self.fontSize['caption'])
    if afterParagraph is not None:
      CCH_DOCX.move_table_after(t,afterParagraph)
    for c in range(cols):
      t.rows[0].cells[c].text = table['title'][c]
    for r,d in enumerate(table['body'],1):
      for c,k in enumerate(table['title']):
        if d[k] is not None:
          t.rows[r].cells[c].text = d[k]
    self.tables += 1
    if caption is None:
      t.rows[rows].cells[0].text = 'Table '+str(self.tables)+'. '+tableName
    else:
      t.rows[rows].cells[0].text = 'Table '+str(self.tables)+'. '+caption
    t.rows[rows].cells[0].merge(t.rows[rows].cells[cols-1])
    return t
    
  @staticmethod
  def parseTags(tag):
    ret = {'type':'text','comment':None,'tagName':None}
    comp = re.compile('<([^<>:]+):(?:([^<>:]+):){0,1}([^<>:]+)>')
    m = comp.match(tag)
    if m is None:
      ret['tagName'] = tag
    else:
      ret['type'] = m.group(1)
      ret['comment'] = m.group(2)
      ret['tagName'] = m.group(3)
    return ret
  
  @staticmethod
  def add_image(image_file,caption=None,run=None,width=None):
    if run is None:
      return
    run.add_text('\r')
    run.add_picture(image_file,width=width)
    run.add_text('\r')
    if caption is not None:
      run.add_text(caption+'\r')
      run.italic = True
      #run.style.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return
    
  @staticmethod
  def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)
  
  @staticmethod
  def set_table_font_size(table, font_size):
    for row in table.rows:
      for cell in row.cells:
          paragraphs = cell.paragraphs
          for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= font_size
                
  @staticmethod
  def _get_heading_level(heading_label):
    return len(heading_label.split('.'))
  
  @staticmethod
  def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

class CCH_HLD_DOCX(CCH_DOCX):
  
  def __init__(self,paramXls=None,contentXls=None,templateFile=None):
    if paramXls is None:
      paramXls = os.path.join(os.path.dirname(__file__),'template','CCH_HLD_Params.xlsx')
    if contentXls is None:
      contentXls = os.path.join(os.path.dirname(__file__),'template','CCH_HLD_Contents.xlsx')
    if templateFile is None:
      templateFile = os.path.join(os.path.dirname(__file__),'template','CCH_HLD_Template.docx')
    super(CCH_HLD_DOCX,self).__init__(templateFile)
    self.loadParameters(paramXls)
    self.loadContents(contentXls)
    self.parseTables()
    self.filterContents()
    self.paragraphs = collections.OrderedDict([
      ('COVER', collections.OrderedDict([('TITLE', None), ('VERSION', None)])),
      ('HISTORY', collections.OrderedDict([('TITLE', None)])),
      ('TOC', collections.OrderedDict([('TITLE', None), ('CONTENT', self.document.paragraphs[0])]))
    ])
    
  def filterContents(self):
    comp = re.compile('(\w+)/(\w+)\s*(==|!=|>|>=|<|<=)\s*(\S+)')
    self.filteredContents = collections.OrderedDict()
    for section in self.contents:
      for p in self.contents[section]['Paragraphs']:
        try:
          m = comp.search(p['Condition'])
          assert m is not None
          k1 = m.group(1)
          k2 = m.group(2)
          op = m.group(3)
          t = m.group(4)
          if op == '==' or '!=':
            cond = eval('"'+str(self.param[k1][k2])+'"'+op+'"'+str(t)+'"')
          else:
            cond = eval(str(self.param[k1][k2])+op+str(t))
          if not cond:
            self.contents[section]['Paragraphs'].remove(p)
        except:
          pass
  
  def getParagraphs(self):
    self.paragraphs = collections.OrderedDict()
    items_p = len(self.document.paragraphs)
    idx_p = 0
    for k in self.paragraphMap:
      for i in range(idx_p, items_p+1):
        p = self.document.paragraphs[i]
        matchPattern = '^\s*'+self.paragraphMap[k]
        if re.search(matchPattern,p.text,re.IGNORECASE) is not None:
          self.paragraphs[k] = p
          idx_p = i+1
          break
  
  def getParagraphVars(self):
    self.paragraphVars = {}
    for k in self.paragraphs:
      p = self.paragraphs[k]
      self.paragraphVars[k] = re.findall('<(\w+)>',p.text)
      
  '''
  def getTables(self):
    self.tables = collections.OrderedDict()
    for n,k in enumerate(self.tableMap):
      self.tables[k] = self.document.tables[n]
  '''
       
  def addDocHistory(self):
    row_cells = self.tables['History'].add_row().cells
    row_cells[0].text = self.param['DocumentVersion']
    row_cells[1].text = ''
    row_cells[2].text = self.param['DocumentAuthor']
    row_cells[3].text = self.param['DocumentDate']
  
  def buildDocFramework(self):
    # Get the cursor paragraph
    pCursor = self.paragraphs['TOC']['CONTENT']
    
    # Cover Page
    p = pCursor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\n'*4+self.param['Generic']['ProjectName']+'\n'+self.param['Generic']['DocumentTitle']+' Document')
    r.font.size = Pt(22)
    r.bold = True
    #p = pCursor.insert_paragraph_before('\n'*4+self.param['ProjectName']+'\n'+self.param['DocumentTitle']+' Document',style=self.document.styles['Title'])
    self.paragraphs['COVER']['TITLE'] = p
    
    p = pCursor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Version '+self.param['Generic']['DocumentVersion']+'\n'*3)
    r.font.size = Pt(16)
    r.bold = True
    self.paragraphs['COVER']['VERSION'] = p
    #p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r.add_picture('template/nokia.png',width=Inches(1.25))
    r.add_break(WD_BREAK.PAGE)
    
    # Document History
    p = pCursor.insert_paragraph_before('Document History',style=self.document.styles['Subtitle'])
    self.addTable('Table-HISTORY', 'Document History', p)
    p = pCursor.insert_paragraph_before()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

    # TOC
    p = pCursor.insert_paragraph_before('Table Of Content',style=self.document.styles['Subtitle'])
    self.paragraphs['TOC']['TITLE'] = p
    r = pCursor.add_run()
    r.add_break(WD_BREAK.PAGE)
    
    ######## Document Body ########
    
    for k in self.contents:
      try:
        self.paragraphs[k]
      except:
        self.paragraphs[k] = {'TITLE':None,'CONTENT':[]}
      h = self.document.add_paragraph(k+'. '+self.contents[k]['Title'],style='Heading '+str(CCH_DOCX._get_heading_level(k)))
      self.paragraphs[k]['TITLE'] = h
      for c in self.contents[k]['Paragraphs']:
        try:
          self.paragraphs[k]['CONTENT']
        except:
          self.paragraphs[k]['CONTENT'] = []
        tokens = []
        if c['Contents'] is None: continue
        m = self.tagre.findall(c['Contents'])
        for l in m:
          for i in range(2):
            try:
              assert len(l[i]) > 0
              tokens.append(l[i])
            except:
              pass
        p = self.document.add_paragraph(style='Normal')
        for token in tokens:
          tag = CCH_DOCX.parseTags(token)
          if tag['type'] == 'param':
            p.add_run(self.applyParam(tag['tagName']))
          elif tag['type'] == 'image':
            r = p.add_run()
            imageName = os.path.join(os.path.dirname(__file__),'template',tag['tagName'])
            self.images += 1
            if tag['comment'] is None:
              caption = 'Figure '+str(self.images)+'.'
            else:
              caption = 'Figure '+str(self.images)+'. '+tag['comment']
            CCH_DOCX.add_image(imageName,caption,r,Inches(5.6))
          elif tag['type'] == 'table':
            self.addTable('Table-'+tag['tagName'], tag['comment'], p)
          else:
            p.add_run(token)
        self.paragraphs[k]['CONTENT'].append(p)
    

      
    